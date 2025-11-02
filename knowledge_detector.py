"""
knowledge_detector_v6.py
Unified Knowledge Detector (v6)
Merges merchant_v5 (RAG, async, doc upload) and admin_merchant_v4 (admin analytics, global reindex)
Provides a single class for all knowledge-based retrieval, indexing, and management.
"""

import asyncio
import logging
import os
import re
import io
from datetime import datetime, timezone
from typing import Dict, List, Optional, Any

from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

logger = logging.getLogger(__name__)

# ======================================
# UNIFIED KNOWLEDGE DETECTOR (v6)
# ======================================
class KnowledgeDetectorV6:
    """
    Unified Knowledge Detector for both Merchant (RAG, indexing)
    and Admin (analytics, global management) use cases.
    """
    def __init__(self, db, chroma_root: str = "./data/chroma_v6"):
        self.db = db  # Database module instance
        self.root_dir = chroma_root
        
        # Initialize embeddings model
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
        
        # In-memory cache for active Chroma collections
        self.collections: Dict[str, Chroma] = {}
        
        # Text splitter for chunking documents
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=500, chunk_overlap=50
        )

        logger.info(f"Initialized Unified KnowledgeDetectorV6 (Merchant + Admin) at {self.root_dir}")

    # ======================================
    # Internal Utility Methods
    # ======================================

    async def run_blocking(self, func, *args, **kwargs):
        """
        Run blocking function in an executor (async-safe).
        Maintained for compatibility, but asyncio.to_thread is preferred.
        """
        try:
            loop = asyncio.get_running_loop()
        except RuntimeError:
            loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, lambda: func(*args, **kwargs))

    def _safe_name(self, identifier: str) -> str:
        """Sanitize an identifier to be a valid collection name."""
        return re.sub(r"[^a-zA-Z0-9_-]", "_", identifier)

    async def _get_collection(self, owner_id: str, is_admin: bool = False) -> Chroma:
        """
        Get or create a persistent Chroma collection for a merchant or admin.
        """
        # Admin collections are prefixed, merchant IDs are sanitized
        name = f"admin_{self._safe_name(owner_id)}" if is_admin else f"merchant_{self._safe_name(owner_id)}"
        
        if name not in self.collections:
            path = os.path.join(self.root_dir, name)
            
            # This will load the collection if it exists, or create it if it doesn't.
            # Running this blocking I/O in a separate thread.
            def _load_chroma():
                return Chroma(
                    collection_name=name,
                    embedding_function=self.embeddings,
                    persist_directory=path
                )
            
            self.collections[name] = await asyncio.to_thread(_load_chroma)
            logger.debug(f"Loaded Chroma collection: {name} from {path}")
            
        return self.collections[name]

    # ======================================
    # Merchant-Focused Methods
    # ======================================

    async def index_product_catalog(self, merchant_id: str):
        """
        Fetch all products for a merchant and index them into their collection.
        (Using richer v5 implementation)
        """
        try:
            products = await self.db.get_all_products(merchant_id)
            if not products:
                logger.warning(f"No products found for merchant {merchant_id}")
                return

            documents = []
            for product in products:
                # Create detailed text representation for embedding
                text = (
                    f"Product: {product['product_name']}\n"
                    f"Price: Rs.{product.get('price', 0):.2f} per {product.get('unit', 'piece')}\n"
                    f"Stock: {product.get('stock', 0)} {product.get('unit', 'pieces')}\n"
                )
                if product.get("description"):
                    text += f"Description: {product['description']}\n"
                if product.get("category"):
                    text += f"Category: {product['category']}\n"
                
                doc = Document(
                    page_content=text,
                    metadata={
                        "product_id": product["product_id"],
                        "product_name": product["product_name"],
                        "price": product.get("price", 0),
                        "stock": product.get("stock", 0),
                        "type": "product",
                        "merchant_id": merchant_id,
                        "indexed_at": datetime.now(timezone.utc).isoformat()
                    }
                )
                documents.append(doc)

            if not documents:
                logger.info(f"No documents generated for {merchant_id} catalog.")
                return

            collection = await self._get_collection(merchant_id)
            await asyncio.to_thread(collection.add_documents, documents)
            await asyncio.to_thread(collection.persist)
            logger.info(f"Indexed {len(documents)} products for {merchant_id}")
            
        except Exception as e:
            logger.error(f"Error indexing product catalog for {merchant_id}: {e}", exc_info=True)

    async def update_context_after_order(self, order_data: Dict):
        """
        Add knowledge about a new order and update product popularity stats.
        (from v5)
        """
        try:
            merchant_id = order_data["merchant_id"]
            collection = await self._get_collection(merchant_id)
            
            # Create text representation of the order
            order_text = (
                f"Order {order_data['order_id']} by customer.\n"
                f"Status: {order_data['status']}\n"
                f"Total: Rs.{order_data.get('total_amount', 0):.2f}\n"
                f"Items: {order_data.get('item_count', 0)}\n"
            )
            for item in order_data.get("items", []):
                order_text += f"- {item['product_name']}: {item['quantity']} {item.get('unit', 'pcs')}\n"
            
            doc = Document(
                page_content=order_text,
                metadata={
                    "order_id": order_data["order_id"],
                    "customer_phone": order_data.get("customer_phone"),
                    "type": "order",
                    "status": order_data["status"],
                    "total_amount": order_data.get("total_amount", 0),
                    "merchant_id": merchant_id,
                    "indexed_at": datetime.now(timezone.utc).isoformat()
                }
            )
            await asyncio.to_thread(collection.add_documents, [doc])

            # Update product metadata stats (runs in parallel)
            stat_tasks = [
                self._update_product_order_stats(
                    merchant_id=merchant_id,
                    product_id=item["product_id"],
                    product_name=item["product_name"]
                ) for item in order_data.get("items", [])
            ]
            await asyncio.gather(*stat_tasks)
            
            await asyncio.to_thread(collection.persist)
            logger.info(f"Updated knowledge base with order {order_data['order_id']}")
            
        except Exception as e:
            logger.error(f"Error updating context after order: {e}", exc_info=True)

    async def _update_product_order_stats(self, merchant_id: str, product_id: str, product_name: str):
        """
        Helper to update product order counts in MongoDB.
        (from v5)
        """
        try:
            # Assuming db.db provides access to raw pymongo collections as in v5
            orders_collection = self.db.db["orders"]
            order_count = await orders_collection.count_documents({
                "merchant_id": merchant_id,
                "items.product_id": product_id
            })
            
            product_metadata = self.db.db["product_metadata"]
            await product_metadata.update_one(
                {"merchant_id": merchant_id, "product_id": product_id},
                {
                    "$set": {
                        "product_name": product_name,
                        "order_count": order_count,
                        "updated_at": datetime.now(timezone.utc)
                    }
                },
                upsert=True
            )
        except Exception as e:
            logger.error(f"Error updating product order stats for {product_id}: {e}", exc_info=True)

    async def add_custom_knowledge(self, merchant_id: str, knowledge_text: str, metadata: Optional[Dict] = None):
        """
        Add arbitrary text to a merchant's knowledge base.
        (from v5, replaces v6 template's add_custom_entry)
        """
        try:
            chunks = self.splitter.split_text(knowledge_text)
            collection = await self._get_collection(merchant_id)
            
            documents = []
            for idx, chunk in enumerate(chunks):
                merged_metadata = dict(metadata) if metadata else {}
                merged_metadata.update({
                    "type": "custom_knowledge",
                    "chunk_index": idx,
                    "merchant_id": merchant_id,
                    "indexed_at": datetime.now(timezone.utc).isoformat()
                })
                doc = Document(page_content=chunk, metadata=merged_metadata)
                documents.append(doc)

            await asyncio.to_thread(collection.add_documents, documents)
            await asyncio.to_thread(collection.persist)
            logger.info(f"Added {len(documents)} custom knowledge chunks for {merchant_id}")
            
        except Exception as e:
            logger.error(f"Error adding custom knowledge for {merchant_id}: {e}", exc_info=True)

    async def search_knowledge(self, merchant_id: str, query: str, k: int = 3) -> Optional[List[str]]:
        """
        Search a merchant's knowledge base for relevant context.
        (from v6 template)
        """
        try:
            collection = await self._get_collection(merchant_id)
            results = await asyncio.to_thread(collection.similarity_search, query, k=k)
            
            if not results:
                return None
            
            return [doc.page_content for doc in results]
        
        except Exception as e:
            logger.error(f"Error searching knowledge for {merchant_id}: {e}", exc_info=True)
            return None

    async def get_popular_products(self, merchant_id: str, limit: int = 5) -> List[Dict]:
        """
        Get most frequently ordered products.
        (from v5)
        """
        try:
            product_metadata = self.db.db["product_metadata"]
            cursor = product_metadata.find(
                {"merchant_id": merchant_id}
            ).sort("order_count", -1).limit(limit)
            
            # Use await cursor.to_list(limit) as in v5
            products = await cursor.to_list(limit)
            return products
        except Exception as e:
            logger.error(f"Error getting popular products for {merchant_id}: {e}", exc_info=True)
            return []

    # ======================================
    # Document Management Methods (from v5)
    # ======================================

    def _extract_pdf(self, content: bytes) -> str:
        """Extract text from PDF bytes."""
        try:
            import PyPDF2
        except ImportError:
            logger.error("PyPDF2 not installed. Run: pip install PyPDF2")
            raise
            
        reader = PyPDF2.PdfReader(io.BytesIO(content))
        text = ""
        for page in reader.pages:
            text += (page.extract_text() or "") + "\n"
        return text

    def _extract_docx(self, content: bytes) -> str:
        """Extract text from DOCX bytes."""
        try:
            import docx
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            raise
            
        doc = docx.Document(io.BytesIO(content))
        return "\n".join([para.text for para in doc.paragraphs if para.text])

    async def upload_document(
        self,
        merchant_id: str,
        file_content: bytes,
        filename: str,
        file_type: str,
        metadata: Optional[Dict] = None
    ) -> Dict:
        """
        Upload and process a document (PDF, DOCX, TXT) into the knowledge base.
        """
        try:
            text = ""
            # Extract text based on file type in a thread
            if file_type == "application/pdf":
                text = await asyncio.to_thread(self._extract_pdf, file_content)
            elif file_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword"
            ]:
                text = await asyncio.to_thread(self._extract_docx, file_content)
            elif file_type.startswith("text/"):
                text = file_content.decode('utf-8')
            else:
                raise ValueError(f"Unsupported file type: {file_type}")

            if not text.strip():
                raise ValueError(f"No text extracted from document: {filename}")

            # Chunk text
            chunks = self.splitter.split_text(text)
            
            # Create LangChain Documents
            documents = []
            doc_id = f"{self._safe_name(merchant_id)}_{self._safe_name(filename)}_{datetime.now(timezone.utc).timestamp()}"
            
            for idx, chunk in enumerate(chunks):
                merged_metadata = dict(metadata or {})
                merged_metadata.update({
                    "doc_id": doc_id,
                    "filename": filename,
                    "chunk_index": idx,
                    "file_type": file_type,
                    "type": "uploaded_document",
                    "merchant_id": merchant_id,
                    "uploaded_at": datetime.now(timezone.utc).isoformat()
                })
                doc = Document(page_content=chunk, metadata=merged_metadata)
                documents.append(doc)
            
            # Add to vector store
            collection = await self._get_collection(merchant_id, is_admin=False)
            await asyncio.to_thread(collection.add_documents, documents)
            await asyncio.to_thread(collection.persist)
            
            # Save metadata record to MongoDB
            doc_record = {
                "merchant_id": merchant_id,
                "doc_id": doc_id,
                "filename": filename,
                "file_type": file_type,
                "chunks_count": len(chunks),
                "uploaded_at": datetime.now(timezone.utc).isoformat(),
                "metadata": metadata or {}
            }
            
            await self.db.db["knowledge_documents"].insert_one(doc_record)
            
            logger.info(f"✅ Uploaded {filename}: {len(chunks)} chunks for merchant {merchant_id}")
            return {
                "status": "success", "doc_id": doc_id,
                "chunks": len(chunks), "filename": filename
            }
        
        except Exception as e:
            logger.error(f"❌ Upload document error for {merchant_id} ({filename}): {e}", exc_info=True)
            return {"status": "error", "message": str(e)}

    async def get_documents(self, merchant_id: str) -> List[Dict]:
        """Get all uploaded document metadata records for a merchant."""
        try:
            docs = await self.db.db["knowledge_documents"].find({
                "merchant_id": merchant_id
            }).sort("uploaded_at", -1).to_list(100)
            
            # Convert ObjectId to string for JSON serialization
            return [{**d, "_id": str(d.get("_id"))} for d in docs]
        
        except Exception as e:
            logger.error(f"❌ Get documents error for {merchant_id}: {e}", exc_info=True)
            return []
    
    async def delete_document(self, merchant_id: str, doc_id: str) -> Dict:
        """
        Delete a document's chunks from vector store and its metadata from DB.
        """
        try:
            # 1. Delete metadata from MongoDB
            result = await self.db.db["knowledge_documents"].delete_one({
                "merchant_id": merchant_id,
                "doc_id": doc_id
            })
            
            if result.deleted_count == 0:
                logger.warning(f"No document metadata found for {doc_id} to delete.")
                # Still try to delete from vector store
            
            # 2. Delete document chunks from Chroma
            collection = await self._get_collection(merchant_id, is_admin=False)
            
            # ChromaDB delete by metadata filter
            await asyncio.to_thread(
                collection.delete,
                where={"doc_id": doc_id}
            )
            await asyncio.to_thread(collection.persist)

            logger.info(f"✅ Deleted document {doc_id} and its chunks for merchant {merchant_id}")
            return {"status": "success", "message": "Document deleted"}
        
        except Exception as e:
            logger.error(f"❌ Delete document error for {doc_id}: {e}", exc_info=True)
            return {"status": "error", "message": str(e)}

    # ======================================
    # Admin-Focused Methods
    # ======================================

    async def global_search(self, query: str, k: int = 5) -> Dict[str, Any]:
        """
        Search across ALL merchant (and admin) collections.
        (from v6 template)
        """
        responses = {}
        # Iterate over a snapshot of collection names
        collection_names = list(self.collections.keys())
        
        for name in collection_names:
            try:
                col = self.collections[name]
                results = await asyncio.to_thread(col.similarity_search, query, k=k)
                if results:
                    responses[name] = [r.page_content for r in results]
            except Exception as e:
                logger.warning(f"Could not search collection {name}: {e}")
                
        return responses

    async def reindex_merchant(self, merchant_id: str):
        """
        Perform a full re-index for a single merchant.
        (Enhanced v6 template using v5 logic)
        """
        logger.info(f"Starting full reindex for merchant {merchant_id}")
        safe_name = f"merchant_{self._safe_name(merchant_id)}"

        # 1. Clear existing collection
        if safe_name in self.collections:
            collection = self.collections.pop(safe_name)
            await asyncio.to_thread(collection.delete_collection)
            logger.info(f"Cleared vector store for merchant {merchant_id}")

        # 2. Re-index products
        await self.index_product_catalog(merchant_id)

        # 3. Re-index recent orders (assuming db method exists)
        try:
            recent_orders = await self.db.get_orders_by_merchant(merchant_id, limit=100)
            for order in recent_orders:
                await self.update_context_after_order(order)
            logger.info(f"Re-indexed {len(recent_orders)} orders for {merchant_id}")
        except Exception as e:
            logger.warning(f"Could not re-index orders for {merchant_id}: {e}")

        # 4. Re-index custom knowledge (assuming db method exists)
        try:
            # Assuming db.db provides access to raw pymongo collections as in v5
            knowledge_entries = await self.db.db["knowledge_base"].find({
                "merchant_id": merchant_id
            }).to_list(1000)
            
            for entry in knowledge_entries:
                knowledge_text = entry.get("content") or entry.get("text")
                if knowledge_text:
                    await self.add_custom_knowledge(
                        merchant_id=merchant_id,
                        knowledge_text=knowledge_text,
                        metadata={"entry_id": str(entry.get("_id"))}
                    )
            logger.info(f"Re-indexed {len(knowledge_entries)} custom knowledge entries for {merchant_id}")
        except Exception as e:
            logger.warning(f"Could not re-index custom knowledge for {merchant_id}: {e}")

        # 5. Re-index uploaded documents
        try:
            doc_records = await self.get_documents(merchant_id)
            # This is complex as it requires re-fetching files.
            # For now, we assume chunks are still in the DB unless deleted.
            # A full re-index would require re-processing files.
            # This implementation assumes re-indexing data sources (DB),
            # not re-uploading files.
            logger.info(f"Found {len(doc_records)} document records. Re-indexing from source files is not yet implemented.")
            # Note: A true re-index of files would require re-running `upload_document`
            # logic for each file, which implies file storage access.
            
        except Exception as e:
            logger.warning(f"Could not re-index documents for {merchant_id}: {e}")
            
        logger.info(f"Completed full reindex for merchant {merchant_id}")

    async def trigger_reindex_all(self):
        """
        Trigger a reindex for ALL known merchants.
        (from v6 template)
        """
        logger.info("Triggering global reindex across all merchants.")
        try:
            # Assuming db method exists
            merchants = await self.db.get_all_merchants() 
            merchant_ids = [m['merchant_id'] for m in merchants if 'merchant_id' in m]
            
            tasks = [self.reindex_merchant(m_id) for m_id in merchant_ids]
            await asyncio.gather(*tasks)
            
            logger.info(f"Global reindex triggered for {len(merchant_ids)} merchants.")
            
        except Exception as e:
            logger.error(f"Failed to trigger global reindex: {e}", exc_info=True)


    # ======================================
    # Knowledge Status & Analytics
    # ======================================

    async def get_knowledge_status(self, owner_id: str, admin: bool = False) -> Dict[str, Any]:
        """
        Get analytics for a specific vector collection.
        (from v6 template)
        """
        doc_count = -1
        last_updated = "N/A"
        path = "N/A"
        
        try:
            collection = await self._get_collection(owner_id, is_admin=admin)
            path = collection._persist_directory
            
            # Run blocking _collection.count() in a thread
            doc_count = await asyncio.to_thread(collection._collection.count)
            
            # Get last modified time of the persist directory
            if os.path.exists(path):
                last_mod_time = os.path.getmtime(path)
                last_updated = datetime.fromtimestamp(last_mod_time, tz=timezone.utc).isoformat()
            
        except Exception as e:
            logger.warning(f"Could not get status for {owner_id}: {e}")
            doc_count = -1
        
        return {
            'owner_id': owner_id,
            'role': 'admin' if admin else 'merchant',
            'doc_count': doc_count,
            'path': path,
            'last_updated': last_updated,
            'status_checked_at': datetime.now(timezone.utc).isoformat()
        }

    def close(self):
        """
        Clean up in-memory collection cache.
        (from v6 template)
        """
        count = len(self.collections)
        for name, col in list(self.collections.items()):
            try:
                # col.persist() # Persist is now often implicit on add/delete
                del self.collections[name]
            except Exception as e:
                logger.warning(f"Error closing collection {name}: {e}")
        
        logger.info(f"Closed {count} Chroma collections from memory cache.")
